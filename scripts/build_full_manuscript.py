from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path('/workspace/scratch/bc83d18af383')
OUT = ROOT / 'paper2_rebuild' / '第十五阶段_完整中文论文合并稿.docx'
ST = ROOT / 'paper2_rebuild'
MAP = ROOT / 'paper2_source' / 'struct_media' / 'word' / 'media' / 'image2.png'
FIG_DIR = ST / 'figures_rebuilt'
FIG2 = FIG_DIR / 'Figure2_depth_harmonization.png'
FIG3 = FIG_DIR / 'Figure3_multiscale_validation.png'
FIG4 = FIG_DIR / 'Figure4_buffer_depth_dependence.png'
FIG5 = FIG_DIR / 'Figure5_robustness_and_explanation.png'

FILES = {
    'intro': ST / '第十二阶段_引言章节重写稿.docx',
    'methods': ST / '第十一阶段_数据与方法章节重写稿.docx',
    'results': ST / '第九阶段_结果章节重写稿.docx',
    'discussion': ST / '第十阶段_讨论章节重写稿.docx',
    'conclusion': ST / '第十三阶段_结论题目与关键词重写稿.docx',
    'abstract': ST / '第十四阶段_中英文摘要重写稿.docx',
}

def set_font(run, east='Noto Serif SC', latin='Times New Roman', size=10.5, bold=None, italic=None):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east)
    run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)

def style_font(style, east, latin, size, bold=False):
    style.font.name = latin
    style._element.rPr.rFonts.set(qn('w:eastAsia'), east)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)

def add_field(paragraph, code):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = code
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t'); text.text = '1'
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    for el in (begin, instr, sep, text, end): run._r.append(el)
    set_font(run, size=9)

def add_text(doc, text, style=None, first_indent=True):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    if style is None:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0.74) if first_indent else None
    return p

def heading_level(text):
    if text and text[0].isdigit() and ' ' in text:
        prefix = text.split()[0]
        return 2 if '.' in prefix else 1
    return None

def copy_section(doc, path, start, stop, skip_prefixes=(), replacements=None, special=None):
    replacements = replacements or {}
    paras = Document(path).paragraphs
    active = False
    for p0 in paras:
        t = p0.text.strip()
        if t == start: active = True
        if not active: continue
        if t.startswith(stop): break
        if not t or any(t.startswith(x) for x in skip_prefixes): continue
        for a,b in replacements.items(): t = t.replace(a,b)
        if special and special(doc, t): continue
        lvl = heading_level(t)
        add_text(doc, t, f'Heading {lvl}' if lvl else None)

def add_figure(doc, path, caption, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8)
    p.add_run().add_picture(str(path), width=Inches(width))
    c = doc.add_paragraph()
    c.style = doc.styles['Caption']
    c.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    c.paragraph_format.keep_together = True
    c.add_run(caption)

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.82)
sec.left_margin = sec.right_margin = Inches(0.9)
sec.header_distance = sec.footer_distance = Inches(0.35)

normal = doc.styles['Normal']
style_font(normal, 'Noto Serif SC', 'Times New Roman', 10.5)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.line_spacing = 1.25
normal.paragraph_format.space_after = Pt(5)
for name, size, before, after in [('Heading 1',16,16,8), ('Heading 2',13,12,6), ('Heading 3',11.5,8,4)]:
    s=doc.styles[name]; style_font(s,'Noto Serif SC','Times New Roman',size,True)
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    s.paragraph_format.keep_with_next=True; s.paragraph_format.keep_together=True
cap=doc.styles['Caption']; style_font(cap,'Noto Serif SC','Times New Roman',9)
cap.paragraph_format.space_before=Pt(3); cap.paragraph_format.space_after=Pt(8); cap.paragraph_format.line_spacing=1.1

title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(5)
r=title.add_run('空间隔离增强下中国土壤有机碳机器学习预测性能的下降及其深度差异')
set_font(r,'Noto Serif SC','Times New Roman',17,True)
sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER; sub.paragraph_format.space_after=Pt(12)
r=sub.add_run('Declining machine-learning performance for soil organic carbon prediction with increasing spatial separation across soil depths in China')
set_font(r,size=11.5,italic=True)

absdoc=Document(FILES['abstract'])
texts=[p.text.strip() for p in absdoc.paragraphs]
for label, body_i, key_i in [('摘要',3,4),('Abstract',6,7)]:
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(5)
    r=p.add_run(label); set_font(r,'Noto Serif SC','Times New Roman',12.5,True)
    p=add_text(doc,texts[body_i],first_indent=False); p.paragraph_format.line_spacing=1.2
    p=add_text(doc,texts[key_i],first_indent=False)
    for run in p.runs: run.bold=True

copy_section(doc, FILES['intro'], '1 引言', '编辑说明',
             skip_prefixes=('第十二阶段重写稿',))

map_done=False
def method_special(doc, t):
    global map_done
    if t == '2.2 标准深度协调' and not map_done:
        add_figure(doc, MAP, '图1 中国SOC分析剖面的空间分布与气候区覆盖。颜色表示按年均温划分的气候区；n为各气候区剖面数。', 6.1)
        map_done=True
    return False
copy_section(doc, FILES['methods'], '2 数据与方法', '编辑说明',
             skip_prefixes=('第十一阶段重写稿', '旧稿中的OLS、随机森林和XGBoost比较'), special=method_special)

fig_done=False
def result_special(doc, t):
    global fig_done
    if t == '3.2 模型性能随空间隔离程度增加而下降':
        add_figure(doc, FIG2, '图2 标准深度协调后的数据保留情况。A，覆盖率不低于50%时各标准深度保留的剖面数；B，不同覆盖率阈值下的剖面保留数量；C，主分析数据中各标准深度的SOC中位数。', 6.5)
    elif t == '3.3 显式空间缓冲揭示了表层与深层SOC泛化能力的差异':
        add_figure(doc, FIG3, '图3 多尺度空间验证揭示模型性能随外推距离增加而下降。A，随机森林和Ridge在不同网格尺度下的性能，虚线表示相应模型的剖面级随机验证结果；B，随机森林性能与测试—训练中位距离的关系；C，不同标准深度在多尺度空间验证下的性能。误差线表示四种网格原点偏移间的标准差。', 6.5)
    elif t == '3.4 容重、土层几何和气候提供了主要预测信息':
        add_figure(doc, FIG4, '图4 显式空间缓冲造成的性能下降及其深度差异。A，实际缓冲与等训练量无缓冲对照；B，不同标准深度的缓冲验证结果；C，200 km条件下选择性移除邻近训练信息造成的附加性能损失。', 6.5)
    return False
copy_section(doc, FILES['results'], '3 结果', '编辑说明',
             skip_prefixes=('第九阶段重写稿','图8  '), replacements={
                 '图8A':'图3A、B', '图8B':'图4A', '图8C':'图4B', '图8D':'图5C',
                 '深层样本的地理代表性更低。':'深层样本的地理代表性更低（图2）。',
                 '同剖面配对分析中的差值为0.194。':'同剖面配对分析中的差值为0.194（图5A）。',
                 '总体R²由0.572降至0.476，下降0.096。':'总体R²由0.572降至0.476，下降0.096（图5B）。',
                 '与旧稿采用单一5°网格得到一个固定差值相比，多尺度结果表明':'多尺度结果进一步表明'
             }, special=result_special)
add_figure(doc, FIG5, '图5 深度差异的稳健性及模型信息来源。A，等样本量、等空间覆盖和同剖面配对条件下的分层性能；B，完整模型与删除容重后模型的性能；C，空间测试折上的分组置换重要性。区间分别表示重复分析的2.5%–97.5%分位范围、网格原点偏移间的标准差或置换分布的2.5%–97.5%分位范围。', 6.5)

copy_section(doc, FILES['discussion'], '4 讨论', '4.5 研究边界', skip_prefixes=('第十阶段重写稿',))
add_text(doc, '4.5 研究边界', 'Heading 2')
add_text(doc, '本研究仍有几项边界。第一，结论建立在China SOCS Database V10的现有样点分布上，低代表性地区的误差可能被总体指标掩盖。第二，显式缓冲提高空间隔离，也改变训练数据的环境覆盖，本文尚未完全分离地理距离与环境外推。第三，100–200 cm层仅有515个协调后剖面，深层估计的不确定性较高。第四，核心稳健性分析以随机森林为主，Ridge提供方向一致的模型族对照，结论仍需在其他算法中检验。第五，当前数据不能逐条确认容重为实测或模型填补，因此含容重结果需与无容重敏感性分析共同解释。这些限制不改变性能随空间隔离增强而下降的主结果，但限定了其向具体地图产品、其他算法和生态过程推广的范围。')
copy_section(doc, FILES['conclusion'], '5 结论', '题目选择说明', replacements={
    '不支持四个标准层严格单调下降':'不支持性能随每一深度层级增加而必然下降'
})

for h, text in [
    ('数据可用性','China SOCS Database V10可通过Zenodo获取（https://doi.org/10.5281/zenodo.17304024）。本文使用的分析文件校验值和数据流审计表将随代码公开。'),
    ('代码可用性','数据预处理、标准深度协调、模型训练、空间验证与图件生成代码将在公开代码仓库提供（仓库地址待作者确认）。'),
    ('作者贡献','[待作者补充]'),('利益冲突','作者声明不存在利益冲突。'),('致谢','[待作者补充]')]:
    add_text(doc,h,'Heading 1'); add_text(doc,text)

add_text(doc,'参考文献','Heading 1')
refs = [
'Chen, Z. et al. A national soil organic carbon density dataset (2010–2024) in China. Scientific Data 12, 1480 (2025). https://doi.org/10.1038/s41597-025-05863-3',
'Hengl, T. et al. SoilGrids250m: Global gridded soil information based on machine learning. PLoS ONE 12, e0169748 (2017). https://doi.org/10.1371/journal.pone.0169748',
'Hicks Pries, C. E. et al. Deep soil organic carbon response to global change. Annual Review of Ecology, Evolution, and Systematics 54, 375–401 (2023). https://doi.org/10.1146/annurev-ecolsys-102320-085332',
'Jobbágy, E. G. & Jackson, R. B. The vertical distribution of soil organic carbon and its relation to climate and vegetation. Ecological Applications 10, 423–436 (2000). https://doi.org/10.1890/1051-0761(2000)010[0423:TVDOSO]2.0.CO;2',
'Liu, F. et al. Mapping high resolution National Soil Information Grids of China. Science Bulletin 67, 328–340 (2022). https://doi.org/10.1016/j.scib.2021.10.013',
'Meyer, H. & Pebesma, E. Machine learning-based global maps of ecological variables and the challenge of assessing them. Nature Communications 13, 2208 (2022). https://doi.org/10.1038/s41467-022-29838-9',
'Ploton, P. et al. Spatial validation reveals poor predictive performance of large-scale ecological mapping models. Nature Communications 11, 4540 (2020). https://doi.org/10.1038/s41467-020-18321-y',
'Roberts, D. R. et al. Cross-validation strategies for data with temporal, spatial, hierarchical, or phylogenetic structure. Ecography 40, 913–929 (2017). https://doi.org/10.1111/ecog.02881',
'Rumpel, C. & Kögel-Knabner, I. Deep soil organic matter—a key but poorly understood component of terrestrial C cycle. Plant and Soil 338, 143–158 (2011). https://doi.org/10.1007/s11104-010-0391-5',
'Wadoux, A. M. J.-C. et al. Spatial cross-validation is not the right way to evaluate map accuracy. Ecological Modelling 457, 109692 (2021). https://doi.org/10.1016/j.ecolmodel.2021.109692']
for ref in refs:
    p=add_text(doc,ref,first_indent=False)
    p.paragraph_format.left_indent=Cm(0.74); p.paragraph_format.first_line_indent=Cm(-0.74); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.05

footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; add_field(footer,'PAGE')

for p in doc.paragraphs:
    for r in p.runs:
        if r.text: set_font(r, 'Noto Serif SC', 'Times New Roman', r.font.size.pt if r.font.size else (10.5), r.bold, r.italic)

doc.core_properties.title='空间隔离增强下中国土壤有机碳机器学习预测性能的下降及其深度差异'
doc.core_properties.subject='完整中文论文合并稿'
doc.core_properties.author=''
doc.save(OUT)
print(OUT)
