"""
Export manuscripts to Word (.docx) format
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re

ROOT = Path(__file__).resolve().parent.parent
EXPORT = ROOT / 'manuscript' / 'export'

def md_to_docx(md_path, docx_path, title):
    """Convert markdown manuscript to Word document."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Read markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_table = False
    table_rows = []
    
    for line in lines:
        line = line.rstrip('\n')
        
        # Skip empty lines
        if not line.strip():
            if in_table and table_rows:
                # Create table
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False
            continue
        
        # Table detection
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            # Skip separator rows
            if all(set(c) <= {'-', ' ', ':'} for c in cells):
                continue
            table_rows.append(cells)
            in_table = True
            continue
        
        if in_table and table_rows:
            _add_table(doc, table_rows)
            table_rows = []
            in_table = False
        
        # Headings
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_heading(line[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=3)
            continue
        
        # Horizontal rule
        if line.strip() == '---':
            doc.add_paragraph('─' * 50)
            continue
        
        # Bold section headers (e.g., "**Random 5-fold cross-validation.**")
        if line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.bold = True
            continue
        
        # Regular paragraph with inline formatting
        p = doc.add_paragraph()
        _add_formatted_text(p, line)
    
    # Handle remaining table
    if in_table and table_rows:
        _add_table(doc, table_rows)
    
    doc.save(docx_path)
    print(f"Saved: {docx_path}")


def _add_table(doc, rows):
    """Add a table to the document."""
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols, style='Table Grid')
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j < n_cols:
                table.rows[i].cells[j].text = cell
                # Bold header row
                if i == 0:
                    for paragraph in table.rows[i].cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    doc.add_paragraph()  # spacing after table


def _add_formatted_text(paragraph, text):
    """Add text with inline bold/italic/code formatting."""
    # Split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Check for italic
            italic_parts = re.split(r'(\*.*?\*)', part)
            for ip in italic_parts:
                if ip.startswith('*') and ip.endswith('*') and len(ip) > 2:
                    run = paragraph.add_run(ip[1:-1])
                    run.italic = True
                elif ip.startswith('`') and ip.endswith('`'):
                    run = paragraph.add_run(ip[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                else:
                    paragraph.add_run(ip)


# Export Chinese version
cn_md = EXPORT / 'manuscript_cn_v2.md'
cn_docx = EXPORT / 'manuscript_cn_v2.docx'
md_to_docx(cn_md, cn_docx, '深度与空间双重约束下中国土壤有机碳机器学习预测的泛化能力评估')

# Export English version
en_md = EXPORT / 'manuscript_en_v2.md'
en_docx = EXPORT / 'manuscript_en_v2.docx'
md_to_docx(en_md, en_docx, 'Evaluating generalization of machine learning for SOC prediction in China')

print("\nDone. Files:")
print(f"  {cn_docx}")
print(f"  {en_docx}")
